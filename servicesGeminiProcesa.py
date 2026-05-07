import asyncio
import random
import json
from models import ResultadoTesis
import google.generativeai as genai
import os
import time
from utilsPdf import generar_pdf_reporte
from models import ResultadoTesis
from dotenv import load_dotenv
import docx

load_dotenv()

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


def extraer_texto_docx(ruta_archivo:str) -> str:
    """Extraer todo el texto en un archivo .word"""
    try:
        doc = docx.Document(ruta_archivo)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        print(f"Error extrayendo texto del documento: {e}")
        return ""


def extraer_info_completa_docx(ruta_archivo:str)->dict:
    """
    Extrae el texto y los metadatos de formato de un archivo word
    """
    try:
        doc = docx.Document(ruta_archivo)
        
        reporte_secciones = []
        for i, section in enumerate(doc.sections):
            reporte_secciones.append({
                "seccion_nro": i + 1,
                "margenes": {
                    "superior": round(section.top_margin.cm, 2),
                    "inferior": round(section.bottom_margin.cm, 2),
                    "izquierdo": round(section.left_margin.cm, 2),
                    "derecho": round(section.right_margin.cm, 2)
                }
            })
        
        estilos_cuerpo = []
        for para in doc.paragraphs:
            if para.text.strip():
                estilos_cuerpo.append({
                    "texto_breve": para.text[:30] + "...", # Referencia para la IA
                    "interlineado": para.paragraph_format.line_spacing,
                    "alineacion": str(para.alignment),
                    "fuente": para.runs[0].font.name if para.runs and para.runs[0].font.name else "Heredado",
                    "tamano": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else "Heredado"
                })
                
        reporte_tablas = []
        for i, table in enumerate(doc.tables):
            # Analizamos la primera celda como muestra del formato de la tabla
            first_cell_para = table.rows[0].cells[0].paragraphs[0]
            reporte_tablas.append({
                "tabla_nro": i + 1,
                "fuente": first_cell_para.runs[0].font.name if first_cell_para.runs else "Heredado",
                "tamano": first_cell_para.runs[0].font.size.pt if first_cell_para.runs else "Heredado"
            })
        
        return {
            "secciones":reporte_secciones,
            "muestreo_cuerpo":estilos_cuerpo[::5],
            "tablas":reporte_tablas,
            "texto_completo": "\n".join([p.text for p in doc.paragraphs])
        }
    except Exception as e:
        print(f"Error extrayendo información del documento: {e}")
        return {"metadatos": {}, "texto": ""}
    

async def extraer_reglas_gemini(ruta_rubrica:str) -> dict:
    """
    Sube el archivo físicamente a la API de Gemini para que lo lea nativamente 
    y construya un JSON estricto con las reglas de evaluación.
    """
    
    diccionario_reglas = {"forma":{},"fondo":{}}
    #archivo_gemini = None
    texto_rubrica = extraer_texto_docx(ruta_rubrica)
    
    if not texto_rubrica:
        return diccionario_reglas
    
    try:
        #archivo_gemini = genai.upload_file(path=ruta_rubrica)
    
        prompt_sistema=f"""
        Actúa como un experto en extracción de datos y estructuración de documentos académicos. Tu tarea es analizar el documento que te envio que es una una guía normativa y convertirlo en un objeto JSON siguiendo estrictamente una jerarquía de "Forma" y "Fondo" por sección.
        INSTRUCCIONES DE SALIDA:
        Formato: Devuelve UNICAMENTE el código JSON, sin formato markdown.
        Estructura Raíz: El JSON debe tener dos ramas principales: "forma" y "fondo".
        Rama "forma": Incluye todas las reglas generales de presentación (Ejm: márgenes, fuentes, interlineado, normas de citación, medios de entrega, etc.).
        Rama "fondo": Desglosa cada sección del documento (Ejm: Carátula, Jurado, Introducción, Referencias, Anexos, etc.) detallando los requisitos de contenido obligatorios para cada una.
        Restricción Crítica: NO incluyas ninguna etiqueta de citación, números de referencia, ni explicaciones adicionales fuera del JSON.
        Idioma: Todo el contenido del JSON debe estar en español.
        """
        
        modelo = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            system_instruction=prompt_sistema,
            generation_config={"temperature":0.3}
        )
        
        respuesta = None
        intentos = 0
        while intentos < 3: 
            try:
                respuesta = await modelo.generate_content_async([
                    #archivo_gemini,
                    #"Analiza el documento adjunto y extrae las reglas en formato JSON."
                    f"DOCUMENTO NORMATIVO:\n\n{texto_rubrica}"
                ])
                break # si hay exito
            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "Quota exceeded" in error_str:
                    intentos+=1
                    print(f"Límite de API alcanzado leyendo reglas. Pausa de 30s... (Intento {intentos}/3)")
                    await asyncio.sleep(30)
                else:
                    raise e
        if not respuesta:
            raise Exception("Se agotaron los 3 intentos para contactar a Gemini (Reglas). La API está muy saturada.")
         
        contenido_crudo = respuesta.text.strip()
        
        #Limpieza de comillas
        if contenido_crudo.startswith("```json"):
            contenido_crudo = contenido_crudo[7:]
        elif contenido_crudo.startswith("```"):
            contenido_crudo = contenido_crudo[3:]
            
        if contenido_crudo.endswith("```"):
            contenido_crudo = contenido_crudo[:-3]
            
        diccionario_reglas = json.loads(contenido_crudo.strip())
        
    except json.JSONDecodeError as e:
        print(f"Error al decodificar el JSON devuelto por Gemini: {e}")
    except Exception as e:
        print(f"Error general en la llamada a Gemini: {e}")
    #finally:
    #    if archivo_gemini:
    #        try:
    #            genai.delete_file(archivo_gemini.name)
    #        except Exception as e:
    #            print(f"No se pudo eliminar el archivo temporal de Gemini: {e}")

    return diccionario_reglas

def aplanar_errores(errores_crudos):
    if not isinstance(errores_crudos, dict):
        return {}
    errores_limpios = {}
    for llave, valor in errores_crudos.items():
        if isinstance(valor, dict):
            # Formateo mejorado para anidaciones (como las de tu imagen)
            textos_anidados = []
            for sub_k, sub_v in valor.items():
                if isinstance(sub_v, dict):
                     # Si la IA anida un nivel más, lo aplanamos a la fuerza
                     sub_v = " - ".join([str(v) for v in sub_v.values()])
                textos_anidados.append(f"{str(sub_k).replace('_', ' ').title()}: {str(sub_v)}")
            errores_limpios[llave] = "\n".join(textos_anidados)
        elif isinstance(valor, list):
            errores_limpios[llave] = ", ".join([str(v) for v in valor])
        else:
            errores_limpios[llave] = str(valor)
    return errores_limpios


async def procesar_documento_gemini(info: dict, prompt: str, reglas:dict, manager, job_id: str, sem: asyncio.Semaphore) -> ResultadoTesis:
    async with sem:
        tiempo_inicio = time.time()
        archivo_gemini = None
        
        # Variables por defecto en caso de que la IA falle al procesar
        titulo_trabajo = f"Documento: {info['nombre']}"
        nombre_alumnos = ["No detectado"]
        nota = 0.0
        errores_forma_detectados = {}
        errores_fondo_detectados = {}
        bases_de_datos = []
        
        try:
            # 1. Subir la tesis del alumno a los servidores de Gemini
            #archivo_gemini = genai.upload_file(path=info['ruta'])
            
            #1. Extrar el texto de la tesis localmente
            info_doc = extraer_info_completa_docx(info['ruta'])
            
            # 2. Construir el prompt estricto integrando el diccionario de reglas
            prompt_sistema = f"""
            Eres un tribunal académico evaluando una tesis de grado. 
            Evalúa el documento adjunto utilizando ESTRICTAMENTE las siguientes reglas extraídas de la rúbrica oficial. 
            Aquí está el contrato de evaluación (JSON):
            {json.dumps(reglas, ensure_ascii=False, indent=3)}
            
            DATOS TÉCNICOS DEL DOCUMENTO (Usa esto para validar la forma):
            
            {json.dumps({
                "secciones": info_doc["secciones"],
                "muestreo_estilos_cuerpo": info_doc["muestreo_cuerpo"],
                "formato_tablas": info_doc["tablas"]
            }, ensure_ascii=False, indent=2)}
            
            GUÍA DE INTERPRETACIÓN TÉCNICA:
            1. Carátula (Sección 1): Los márgenes y fuentes pueden variar ligeramente. Sé flexible aquí.
            2. Cuerpo (Sección 2 en adelante): Los márgenes DEBEN cumplir estrictamente la regla (Ej: 3cm Izq, 2.5cm otros).
            3. Tablas: Los datos técnicos de 'formato_tablas' indican si el alumno redujo la fuente en tablas. Según la normativa, esto es PERMITIDO. No lo marques como error.
            4. Estilos: Valida que el 'muestreo_estilos_cuerpo' mantenga consistencia con el tipo de fuente y el interlineado.
            
            El trabaj empieza en 20 luego por los errores se va descontando, recuerda los errores de fondo pesan mas que los de forma.
            Instrucciones adicionales del evaluador principal (usuario): {prompt}
            
            INSTRUCCIONES DE SALIDA:
            Debes devolver UNICAMENTE un objeto JSON válido con la siguiente estructura:
            {{
                "titulo_trabajo": "Título de la tesis encontrado en la carátula o inicio",
                "nombre_alumnos": ["Nombre 1", "Nombre 2 (si hay)"],
                "nota": [Calificación numérica de 0 a 20 basada en la cantidad y gravedad de los errores],
                "errores_forma": {{
                    "llave_exacta_de_regla": "EXPLICACIÓN DETALLADA. Describe qué encontraste, en qué página/sección está el error, y cómo el alumno debe corregirlo según la rúbrica. (Mínimo 3 oraciones)."
                }},
                "errores_fondo": {{
                    "llave_exacta_de_regla": "EXPLICACIÓN DETALLADA. Menciona el hallazgo, el impacto en la investigación y la acción correctiva. (Mínimo 3 oraciones)."
                }},
                "bases_de_datos": ["Scopus", "IEEE", "PubMed", "Science Direct"] // Extrae las bases de datos bibliográficas detectadas en el documento.
            }}
            RESTRICCIÓN CRÍTICA DE FORMATO:
            - Prohibido anidar diccionarios dentro de 'errores_forma' o 'errores_fondo'. 
            - Cada valor debe ser un texto plano y extenso (string), NUNCA un sub-objeto o lista.
            - Usa exactamente las mismas llaves del contrato.
            - No incluyas markdown (```json).
            """

            modelo = genai.GenerativeModel(
                model_name="gemini-2.5-flash",
                system_instruction=prompt_sistema,
                generation_config={"temperature": 0.2}
            )
            
            respuesta = None
            intentos = 0
            while intentos < 3: 
                try:
                    # 3. Llamada a la IA
                    respuesta = await modelo.generate_content_async([
                        #archivo_gemini, 
                        #"Evalúa la tesis adjunta detalladamente y devuelve el JSON solicitado."
                        f"TEXTO COMPLETO DE LA TESIS PARA EVALUAR FONDO:\n\n{info_doc['texto_completo']}"
                    ])
                    break # si hay exito
                except Exception as e:
                    error_str = str(e)
                    if "429" in error_str or "Quota exceeded" in error_str:
                        intentos+=1
                        print(f"Límite de API alcanzado leyendo Trabajos. Pausa de 30s... (Intento {intentos}/3)")
                        await asyncio.sleep(30)
                    else:
                        raise e
            
            if not respuesta:
                raise Exception("Se agotaron los 3 intentos para contactar a Gemini (Reglas). La API está muy saturada.")
            
            contenido_crudo = respuesta.text.strip()
            
            # 4. Limpieza de etiquetas Markdown
            if contenido_crudo.startswith("```json"):
                contenido_crudo = contenido_crudo[7:]
            elif contenido_crudo.startswith("```"):
                contenido_crudo = contenido_crudo[3:]
            if contenido_crudo.endswith("```"):
                contenido_crudo = contenido_crudo[:-3]
                
            # 5. Convertir a diccionario de Python
            datos_extraidos = json.loads(contenido_crudo.strip())
            
            # 6. Asignar de forma segura usando .get()
            titulo_trabajo = datos_extraidos.get("titulo_trabajo", titulo_trabajo)
            nombre_alumnos = datos_extraidos.get("nombre_alumnos", nombre_alumnos)
            nota_bruta = float(datos_extraidos.get("nota", 0.0))
            if isinstance(nota_bruta, list) and len(nota_bruta) > 0:
                nota = float(nota_bruta[0])
            else:
                try:
                    nota = float(nota_bruta)
                except:
                    nota = 0.0
            errores_forma_detectados = aplanar_errores(datos_extraidos.get("errores_forma", {}))
            errores_fondo_detectados = aplanar_errores(datos_extraidos.get("errores_fondo", {}))
            bases_de_datos = datos_extraidos.get("bases_de_datos", [])

        except json.JSONDecodeError as e:
            print(f"Error decodificando el JSON de la evaluación de {info['nombre']}: {e}")
            errores_forma_detectados = {"error_sistema": "El modelo no devolvió un formato válido."}
        except Exception as e:
            print(f"Error general evaluando {info['nombre']} con Gemini: {e}")
            errores_forma_detectados = {"error_procesamiento": str(e)}
        finally:
            # 7. IMPORTANTE: Borrar el archivo de la nube de Gemini
            if archivo_gemini:
                try:
                    genai.delete_file(archivo_gemini.name)
                except Exception as e:
                    print(f"No se pudo eliminar la tesis de la nube de Gemini: {e}")
        
        tiempo_fin = time.time()
        
        
        #Procesar PDF
        datos_para_pdf = {
            "nombre_archivo": info['nombre'],
            "titulo_trabajo": titulo_trabajo,
            "nombre_alumnos": nombre_alumnos,
            "nota": nota,
            "estado":"Aprobado" if nota >= 14 else "Desaprobado",
            "errores_forma": errores_forma_detectados,
            "errores_fondo": errores_fondo_detectados,
            "bases_de_datos": bases_de_datos
        }
        
        nombre_pdf = generar_pdf_reporte(datos_para_pdf, job_id)
        
        url_pdf = f"http://127.0.0.1:8000/reportes/{nombre_pdf}"
        
        # 8. Retornar la estructura Pydantic que espera el orquestador y el dashboard
        return ResultadoTesis(
            nombre_archivo=info['nombre'],
            titulo_trabajo=titulo_trabajo,
            nombre_alumnos=nombre_alumnos,
            nota=nota,
            errores_forma=errores_forma_detectados,
            errores_fondo=errores_fondo_detectados,
            bases_de_datos=bases_de_datos,
            tiempo_procesamiento_segundos=round(tiempo_fin - tiempo_inicio, 2),
            link_pdf_review=url_pdf
        )