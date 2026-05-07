import asyncio
import json
import os
import time
from openai import AsyncOpenAI, RateLimitError
import docx
from dotenv import load_dotenv
from utilsPdf import generar_pdf_reporte
from models import ResultadoTesis

# Cargar variables de entorno
load_dotenv()

client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extraer_texto_docx(ruta_archivo: str) -> str:
    """Extraer todo el texto en un archivo .word"""
    try:
        doc = docx.Document(ruta_archivo)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    except Exception as e:
        print(f"Error extrayendo texto del documento: {e}")
        return ""

def extraer_info_completa_docx(ruta_archivo: str) -> dict:
    """Extrae el texto y los metadatos de formato de un archivo word"""
    try:
        doc = docx.Document(ruta_archivo)
        reporte_secciones = []
        for i, section in enumerate(doc.sections):
            reporte_secciones.append({
                "seccion_nro": i + 1,
                "margenes": {
                    "superior": round(section.top_margin.cm, 2),
                    "inferior": round(section.bottom_margin.cm, 2),
                    "izquierdo": round(section.left_margin.cm, 2),
                    "derecho": round(section.right_margin.cm, 2)
                }
            })
        
        estilos_cuerpo = []
        for para in doc.paragraphs:
            if para.text.strip():
                estilos_cuerpo.append({
                    "texto_breve": para.text[:30] + "...", 
                    "interlineado": para.paragraph_format.line_spacing,
                    "alineacion": str(para.alignment),
                    "fuente": para.runs[0].font.name if para.runs and para.runs[0].font.name else "Heredado",
                    "tamano": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else "Heredado"
                })
                
        reporte_tablas = []
        for i, table in enumerate(doc.tables):
            try:
                first_cell_para = table.rows[0].cells[0].paragraphs[0]
                reporte_tablas.append({
                    "tabla_nro": i + 1,
                    "fuente": first_cell_para.runs[0].font.name if first_cell_para.runs else "Heredado",
                    "tamano": first_cell_para.runs[0].font.size.pt if first_cell_para.runs else "Heredado"
                })
            except: continue
        
        return {
            "secciones": reporte_secciones,
            "muestreo_cuerpo": estilos_cuerpo[::5],
            "tablas": reporte_tablas,
            "texto_completo": "\n".join([p.text for p in doc.paragraphs])
        }
    except Exception as e:
        print(f"Error extrayendo información del documento: {e}")
        return {"metadatos": {}, "texto": ""}

def aplanar_errores(errores_crudos):
    if not isinstance(errores_crudos, dict): return {}
    errores_limpios = {}
    for llave, valor in errores_crudos.items():
        if isinstance(valor, dict):
            textos_anidados = []
            for sub_k, sub_v in valor.items():
                if isinstance(sub_v, dict):
                     sub_v = " - ".join([str(v) for v in sub_v.values()])
                textos_anidados.append(f"{str(sub_k).replace('_', ' ').title()}: {str(sub_v)}")
            errores_limpios[llave] = "\n".join(textos_anidados)
        elif isinstance(valor, list):
            errores_limpios[llave] = ", ".join([str(v) for v in valor])
        else:
            errores_limpios[llave] = str(valor)
    return errores_limpios

async def extraer_reglas_openai(ruta_rubrica: str) -> dict:
    diccionario_reglas = {"forma": {}, "fondo": {}}
    texto_rubrica = extraer_texto_docx(ruta_rubrica)
    
    if not texto_rubrica: return diccionario_reglas
    
    prompt_sistema = """
    Actúa como un experto en extracción de datos y estructuración de documentos académicos. 
    Analiza la guía normativa y conviértela en un objeto JSON estricto.
    Estructura Raíz: "forma" y "fondo".
    Rama "forma": Reglas generales de presentación (márgenes, fuentes, etc).
    Rama "fondo": Requisitos de contenido obligatorios por sección.
    Idioma: Español.
    """
    
    intentos = 0
    while intentos < 3: 
        try:
            respuesta = await client.chat.completions.create(
                model="gpt-4o-mini",
                response_format={ "type": "json_object" }, 
                messages=[
                    {"role": "system", "content": prompt_sistema},
                    {"role": "user", "content": f"DOCUMENTO NORMATIVO:\n\n{texto_rubrica}"}
                ],
                temperature=0.3
            )
            
            contenido_crudo = respuesta.choices[0].message.content
            diccionario_reglas = json.loads(contenido_crudo)
            break
            
        except RateLimitError:
            intentos += 1
            print(f"Límite de API OpenAI alcanzado (Reglas). Pausa 20s... (Intento {intentos}/3)")
            await asyncio.sleep(20)
        except Exception as e:
            print(f"Error general OpenAI: {e}")
            break

    return diccionario_reglas


async def procesar_documento_openai(info: dict, prompt: str, reglas: dict, manager, job_id: str, sem: asyncio.Semaphore) -> ResultadoTesis:
    async with sem:
        tiempo_inicio = time.time()
        
        titulo_trabajo = f"Documento: {info['nombre']}"
        nombre_alumnos = ["No detectado"]
        nota = 0.0
        errores_forma_detectados = {}
        errores_fondo_detectados = {}
        bases_de_datos = []
        
        try:
            info_doc = extraer_info_completa_docx(info['ruta'])
            
            # --- PROMPT MEJORADO PARA DESCRIPCIONES ABUNDANTES ---
            prompt_sistema = f"""
            Eres un tribunal académico de grado sumamente riguroso y analítico.
            Evalúa el documento utilizando ESTRICTAMENTE este contrato de reglas (JSON):
            {json.dumps(reglas, ensure_ascii=False, indent=2)}
            
            DATOS TÉCNICOS DEL DOCUMENTO (Para validar formato):
            {json.dumps({
                "secciones": info_doc["secciones"],
                "muestreo_estilos_cuerpo": info_doc["muestreo_cuerpo"],
                "formato_tablas": info_doc["tablas"]
            }, ensure_ascii=False, indent=2)}
            
            GUÍA TÉCNICA Y PENALIZACIÓN:
            1. Empiezas con 20 puntos. Resta puntos según la gravedad: fondo pesa más que forma.
            2. Carátula: Flexible en márgenes. Cuerpo: Márgenes estrictos.
            3. Tablas: Fuente reducida permitida.
            
            Instrucciones del usuario: {prompt}
            
            INSTRUCCIONES CRÍTICAS DE SALIDA (FORMATO JSON):
            Debes devolver UNICAMENTE un JSON con esta estructura exacta:
            {{
                "titulo_trabajo": "Título de la tesis",
                "nombre_alumnos": ["Nombre 1"],
                "nota": 15,
                "errores_forma": {{ 
                    "llave_exacta_de_regla_1": "string con descripción detallada como se indicaba arriba. Si la regla se cumplió o no, igual debes escribir una descripción detallada explicando por qué se cumplió o no y resaltando las partes del documento que evidencian ese cumplimiento o no.",
                    "llave_exacta_de_regla_2": "string con descripción detallada como se indicaba arriba. Si la regla se cumplió o no, igual debes escribir una descripción detallada explicando por qué se cumplió o no y resaltando las partes del documento que evidencian ese cumplimiento o no.",
                    ...
                    
                }},
                "errores_fondo": {{ 
                    "llave_exacta_de_regla_1": "string con descripción detallada como se indicaba arriba. Si la regla se cumplió o no, igual debes escribir una descripción detallada explicando por qué se cumplió o no y resaltando las partes del documento que evidencian ese cumplimiento o no.",
                    "llave_exacta_de_regla_2": "string con descripción detallada como se indicaba arriba. Si la regla se cumplió o no, igual debes escribir una descripción detallada explicando por qué se cumplió o no y resaltando las partes del documento que evidencian ese cumplimiento o no.",
                    ...
                }},
                "bases_de_datos": ["Scopus", "IEEE", "Scielo","Web of Science","Google Scholar","ProQuest","Redalyc","Esmeral","Otros"]
            }}
            
            IMPORTANTE: 
            - Prohibido anidar diccionarios.
            - Esa "nota":15 es solo un ejemplo, la nota real debe basarse en el análisis riguroso del documento contra las reglas.
            - OBLIGATORIO: Debes iterar sobre TODAS las reglas definidas en el contrato original. Tu salida en 'errores_forma' y 'errores_fondo' debe contener tantas claves como reglas existan en el contrato.
            - En los errores de forma y fondo debes incluir tantas claves como reglas infringidas o aprobadas (la idea es que por cada clave des una descripcion detallada, si esta bien explicando porque y si esta mal explicando igual). Tu descripción detallada debe contener: Regla invocada, Hallazgo (cita la parte donde ocurrio ese problema), Evaluación (aprobado o no), Acción requerida (recomendación de mejora o correción).
            - Esos ejemplos de nombres de "llave_exacta_de_regla_1", "llave_exacta_de_regla_2", etc. son solo ilustrativos. Las llaves deben ser exactamente las mismas que las reglas del JSON de entrada. Si una regla se llama "margen_superior", entonces la llave en errores_forma o errores_fondo debe ser exactamente "margen_superior" para esa regla, y así con cada regla. Esto es fundamental para que el sistema pueda mapear correctamente los errores detectados con las reglas establecidas.
            """

            intentos = 0
            while intentos < 3: 
                try:
                    respuesta = await client.chat.completions.create(
                        model="gpt-4o-mini",
                        response_format={ "type": "json_object" }, 
                        messages=[
                            {"role": "system", "content": prompt_sistema},
                            {"role": "user", "content": f"TEXTO COMPLETO DE LA TESIS:\n\n{info_doc['texto_completo']}"}
                        ],
                        temperature=0.3 # Subimos un poco la temperatura para que sea más locuaz y descriptivo
                    )
                    
                    contenido_crudo = respuesta.choices[0].message.content
                    datos_extraidos = json.loads(contenido_crudo)
                    break 
                    
                except RateLimitError:
                    intentos += 1
                    print(f"Límite API OpenAI (Tesis). Pausa 20s... (Intento {intentos}/3)")
                    await manager.send_personal_message(json.dumps({
                        "estado":"PROCESANDO",
                        "procesados": 0, "total": 0,
                        "mensaje":f"Pausando 20s por límites de API en {info['nombre']}..."
                    }), job_id)
                    await asyncio.sleep(20)
                except Exception as e:
                    raise e
            
            titulo_trabajo = datos_extraidos.get("titulo_trabajo", titulo_trabajo)
            nombre_alumnos = datos_extraidos.get("nombre_alumnos", nombre_alumnos)
            
            nota_bruta = datos_extraidos.get("nota", 0.0)
            if isinstance(nota_bruta, list) and len(nota_bruta) > 0: nota = float(nota_bruta[0])
            else:
                try: nota = float(nota_bruta)
                except: nota = 0.0
                
            errores_forma_detectados = aplanar_errores(datos_extraidos.get("errores_forma", {}))
            errores_fondo_detectados = aplanar_errores(datos_extraidos.get("errores_fondo", {}))
            bases_de_datos = datos_extraidos.get("bases_de_datos", [])

        except Exception as e:
            print(f"Error general evaluando {info['nombre']} con OpenAI: {e}")
            errores_forma_detectados = {"error_procesamiento": str(e)}
        
        tiempo_fin = time.time()
        
        # Procesar PDF
        datos_para_pdf = {
            "nombre_archivo": info['nombre'],
            "titulo_trabajo": titulo_trabajo,
            "nombre_alumnos": nombre_alumnos,
            "nota": nota,
            "estado": "Aprobado" if nota >= 14 else "Desaprobado",
            "errores_forma": errores_forma_detectados,
            "errores_fondo": errores_fondo_detectados,
            "bases_de_datos": bases_de_datos
        }
        
        nombre_pdf = generar_pdf_reporte(datos_para_pdf, job_id)
        url_pdf = f"http://127.0.0.1:8000/reportes/{nombre_pdf}"
        
        return ResultadoTesis(
            nombre_archivo=info['nombre'],
            titulo_trabajo=titulo_trabajo,
            nombre_alumnos=nombre_alumnos,
            nota=nota,
            errores_forma=errores_forma_detectados,
            errores_fondo=errores_fondo_detectados,
            bases_de_datos=bases_de_datos,
            tiempo_procesamiento_segundos=round(tiempo_fin - tiempo_inicio, 2),
            link_pdf_review=url_pdf
        )